#!/usr/bin/env python3
"""Planificador de notas para los slots detectados en InDesign.

El objetivo del script es producir un **reporte previo** con cálculos de
ocupación que indiquen si las notas (título + cuerpo + imagen opcional) entran
en los slots dibujados en el documento.  El reporte sirve como insumo para el
script JSX encargado de volcar el contenido final sobre la maqueta.

Características principales:

* Lee la geometría de los slots a partir de un JSON generado con
  ``detect_slots.jsx`` o, en su defecto, directamente desde un ``.idml``.
* Procesa las notas desde una carpeta con ``.docx`` (detectando títulos en
  negrita) o desde un CSV/JSON de resumen (como ``reporte_cierre``).
* Calcula el alto requerido para cada nota considerando tamaños de título y
  cuerpo configurables, tolerancias y la reserva para fotografías de 2
  columnas.
* Genera un JSON y, opcionalmente, un CSV con el detalle de qué notas entran,
  cuáles quedan fuera y el excedente aproximado en caracteres.

La idea es que el resultado le permita al ``.jsx`` ejecutar directamente el
layout elegido sin volver a hacer la parte de mediciones pesadas.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import os
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from typing import Dict, Iterable, Iterator, List, Optional, Sequence, Tuple

try:
    from docx import Document  # type: ignore
    from docx.document import Document as DocumentType  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore

    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover - import opcional
    Document = None  # type: ignore
    DocumentType = None  # type: ignore
    Paragraph = None  # type: ignore
    DOCX_AVAILABLE = False

try:
    from generate_slot_report import analyze_document

    IDML_AVAILABLE = True
except Exception:  # pragma: no cover - import opcional
    analyze_document = None  # type: ignore
    IDML_AVAILABLE = False


MM_PER_POINT = 25.4 / 72.0
POINTS_PER_CM = 28.3464566929

IMAGE_DEFAULT_WIDTH_CM = 10.145
IMAGE_MIN_HEIGHT_CM = 5.35


@dataclass
class Slot:
    """Información geométrica del slot disponible para una o más notas."""

    page_name: str
    page_index: int
    slot_id: str
    x_pt: float
    y_pt: float
    width_pt: float
    height_pt: float

    def area(self) -> float:
        return self.width_pt * self.height_pt

    def to_dict(self) -> Dict[str, object]:
        return {
            "page_name": self.page_name,
            "page_index": self.page_index,
            "slot_id": self.slot_id,
            "x_pt": self.x_pt,
            "y_pt": self.y_pt,
            "width_pt": self.width_pt,
            "height_pt": self.height_pt,
            "width_mm": self.width_pt * MM_PER_POINT,
            "height_mm": self.height_pt * MM_PER_POINT,
        }


@dataclass
class Note:
    """Representa una nota individual con título, cuerpo e imagen opcional."""

    page_hint: str
    docx_file: Optional[str]
    note_index: int
    title: str
    body_text: str
    title_chars: int
    body_chars: int
    image_count: int = 0

    def has_image(self) -> bool:
        return self.image_count > 0

    def note_id(self) -> str:
        base = self.docx_file or self.title
        return f"{self.page_hint}#{self.note_index}:{base}"


@dataclass
class LayoutResult:
    note: Note
    slot: Slot
    fit: bool
    font_title: float
    font_body: float
    heights_pt: Dict[str, float]
    remaining_height_pt: float
    overflow_chars: int
    used_height_pt: float
    attempt_summary: Dict[str, object] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, object]:
        return {
            "note_id": self.note.note_id(),
            "page_hint": self.note.page_hint,
            "slot_id": self.slot.slot_id,
            "fit": self.fit,
            "font_title_pt": self.font_title,
            "font_body_pt": self.font_body,
            "title_chars": self.note.title_chars,
            "body_chars": self.note.body_chars,
            "image_count": self.note.image_count,
            "heights_pt": self.heights_pt,
            "remaining_height_pt": self.remaining_height_pt,
            "overflow_chars": self.overflow_chars,
            "used_height_pt": self.used_height_pt,
            "title": self.note.title,
            "docx_file": self.note.docx_file,
            "attempt": self.attempt_summary,
        }


@dataclass
class UnplacedNote:
    note: Note
    reason: str
    overflow_chars: int = 0
    missing_height_pt: float = 0.0

    def to_dict(self) -> Dict[str, object]:
        return {
            "note_id": self.note.note_id(),
            "page_hint": self.note.page_hint,
            "reason": self.reason,
            "overflow_chars": self.overflow_chars,
            "missing_height_pt": self.missing_height_pt,
            "title": self.note.title,
            "docx_file": self.note.docx_file,
        }


@dataclass
class ColumnModel:
    column_width_pt: float
    gutter_pt: float
    char_width_factor: float
    leading_factor: float

    def chars_per_line(self, width_pt: float, font_size_pt: float) -> float:
        width_factor = max(width_pt, 1.0)
        return width_factor / (font_size_pt * self.char_width_factor)

    def line_height(self, font_size_pt: float) -> float:
        return font_size_pt * self.leading_factor


def parse_arguments(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Planificador de ocupación de notas")
    parser.add_argument("--slots-json", help="JSON generado con detect_slots.jsx")
    parser.add_argument("--idml", help="Documento IDML a analizar", default=None)
    parser.add_argument("--notes-root", help="Carpeta que contiene subcarpetas por página con DOCX", default=None)
    parser.add_argument(
        "--notes-report",
        help="Resumen CSV/JSON con columnas page_folder, docx_file, note_index, title, char_count, images_count",
        default=None,
    )
    parser.add_argument("--output-json", help="Archivo JSON de salida", default="layout_plan.json")
    parser.add_argument("--output-csv", help="Archivo CSV de salida", default=None)
    parser.add_argument("--pretty", action="store_true", help="Imprime resumen legible en consola")
    parser.add_argument(
        "--config",
        help="Archivo JSON opcional con overrides para tamaños, tolerancias y heurísticas",
        default=None,
    )
    parser.add_argument(
        "--default-image-assignment",
        choices=["balance", "first", "none"],
        default="balance",
        help="Estrategia para asignar imágenes cuando no se detectan automáticamente",
    )
    return parser.parse_args(argv)


def load_config(path: Optional[str]) -> Dict[str, object]:
    base_config: Dict[str, object] = {
        "title": {
            "pt_base": 25.0,
            "pt_min": 24.0,
            "pt_max": 26.0,
            "leading_factor": 1.08,
            "char_width_factor": 0.36,
            "spacing_after_pt": 8.0,
        },
        "body": {
            "pt_base": 9.5,
            "pt_min": 9.0,
            "pt_max": 10.0,
            "leading_factor": 1.2,
            "char_width_factor": 0.42,
            "spacing_after_pt": 6.0,
        },
        "note_spacing_pt": 10.0,
        "image": {
            "width_cm": IMAGE_DEFAULT_WIDTH_CM,
            "min_height_cm": IMAGE_MIN_HEIGHT_CM,
            "spacing_pt": 6.0,
        },
        "column": {
            "gutter_pt": 12.0,
        },
    }

    if not path:
        return base_config

    try:
        with open(path, "r", encoding="utf-8") as fh:
            user_cfg = json.load(fh)
    except Exception as exc:  # pylint: disable=broad-except
        print(f"No se pudo leer config '{path}': {exc}", file=sys.stderr)
        return base_config

    def deep_merge(target: Dict[str, object], source: Dict[str, object]) -> Dict[str, object]:
        for key, value in source.items():
            if isinstance(value, dict) and isinstance(target.get(key), dict):
                target[key] = deep_merge(target[key], value)  # type: ignore[assignment]
            else:
                target[key] = value
        return target

    return deep_merge(base_config, user_cfg)


def load_slots_from_json(path: str) -> List[Slot]:
    with open(path, "r", encoding="utf-8") as fh:
        payload = json.load(fh)

    slots: List[Slot] = []
    if isinstance(payload, dict) and "pages" in payload:
        # Formato del generate_slot_report
        fallback_candidates: List[Slot] = []
        for page in payload.get("pages", []):
            page_name = str(page.get("page_name", page.get("page_index", "")))
            page_index = int(page.get("page_index", 0))
            page_slots: List[Slot] = []
            for idx, item in enumerate(page.get("items", [])):
                bounds = item.get("bounds_page") or item.get("bounds")
                if not bounds:
                    continue
                width_pt = float(bounds.get("width_pt", bounds.get("width", 0)))
                height_pt = float(bounds.get("height_pt", bounds.get("height", 0)))
                x_pt = float(bounds.get("left_pt", bounds.get("x_pt", 0)))
                y_pt = float(bounds.get("top_pt", bounds.get("y_pt", 0)))
                label = str(item.get("label", ""))
                obj_style = str(item.get("object_style", ""))
                slot_id = f"page{page_index}_slot{idx+1}"
                slot = Slot(
                    page_name=page_name,
                    page_index=page_index,
                    slot_id=slot_id,
                    x_pt=x_pt,
                    y_pt=y_pt,
                    width_pt=width_pt,
                    height_pt=height_pt,
                )
                fallback_candidates.append(slot)
                if label and (label.lower().startswith("slot") or label.lower() == "root"):
                    page_slots.append(slot)
                elif obj_style and "slot" in obj_style.lower():
                    page_slots.append(slot)
            if page_slots:
                slots.extend(page_slots)
        if slots:
            return slots
        # Sin etiquetas: usamos heurística de área para filtrar candidatos grandes
        for slot in fallback_candidates:
            width_mm = slot.width_pt * MM_PER_POINT
            height_mm = slot.height_pt * MM_PER_POINT
            if width_mm >= 40 and height_mm >= 40:
                slots.append(slot)
        return slots

    if isinstance(payload, list):
        # Formato plano del detect_slots.jsx
        for idx, item in enumerate(payload):
            page_name = str(item.get("page", ""))
            page_index = int(item.get("page_index", item.get("page", 0))) if str(item.get("page", "")).isdigit() else idx + 1
            x_pt = float(item.get("x_pt", 0))
            y_pt = float(item.get("y_pt", 0))
            width_pt = float(item.get("w_pt", item.get("width_pt", 0)))
            height_pt = float(item.get("h_pt", item.get("height_pt", 0)))
            slot_id = str(item.get("id", f"slot_{idx+1}"))
            slots.append(
                Slot(
                    page_name=page_name,
                    page_index=page_index,
                    slot_id=slot_id,
                    x_pt=x_pt,
                    y_pt=y_pt,
                    width_pt=width_pt,
                    height_pt=height_pt,
                )
            )
        return slots

    raise ValueError("Formato de slots desconocido; se esperaba lista o diccionario con 'pages'")


def load_slots_from_idml(path: str) -> List[Slot]:
    if not IDML_AVAILABLE or analyze_document is None:  # pragma: no cover - fallback
        raise RuntimeError("Soporte IDML no disponible; instale dependencias o use --slots-json")

    document_info, pages, items = analyze_document(path)  # type: ignore[misc]
    slots: List[Slot] = []
    slot_counter = Counter()
    for item in items:
        label = (item.label or "").lower()
        obj_style = (item.object_style or "").lower()
        if "slot" not in label and label != "root" and "slot" not in obj_style:
            continue
        bounds = item.bounds_page or item.bounds_spread
        if not bounds:
            continue
        page_name = ""
        page_index = 0
        if item.page_name:
            page_name = str(item.page_name)
        if item.page_index:
            page_index = int(item.page_index)
        slot_counter[page_name] += 1
        slot_id = f"page{page_index or page_name}_slot{slot_counter[page_name]}"
        slots.append(
            Slot(
                page_name=page_name or str(page_index),
                page_index=page_index,
                slot_id=slot_id,
                x_pt=float(bounds.left),
                y_pt=float(bounds.top),
                width_pt=float(bounds.width),
                height_pt=float(bounds.height),
            )
        )

    if not slots:
        raise RuntimeError(
            "No se detectaron slots etiquetados en el IDML. Ejecutá detect_slots.jsx o etiquetá los marcos con 'slot'."
        )

    return slots


def normalize_page_key(value: str) -> str:
    value = value.strip()
    if not value:
        return value
    return re.sub(r"\s+", " ", value)


def first_int_in_string(value: str) -> Optional[int]:
    match = re.search(r"\d+", value)
    if not match:
        return None
    try:
        return int(match.group(0))
    except ValueError:
        return None


def assign_images(notes: List[Note], available_images: List[str], strategy: str) -> None:
    if not notes:
        return
    if strategy == "none":
        return
    if strategy == "balance":
        per_note = len(available_images) // len(notes)
        remainder = len(available_images) % len(notes)
        for idx, note in enumerate(notes):
            count = per_note + (1 if idx < remainder else 0)
            note.image_count = max(note.image_count, count)
        return
    if strategy == "first":
        for idx, note in enumerate(notes):
            if idx < len(available_images):
                note.image_count = max(note.image_count, 1)
            else:
                break


def is_paragraph_bold(paragraph: Paragraph) -> bool:  # type: ignore[valid-type]
    text = paragraph.text or ""
    if not text.strip():
        return False
    total_chars = 0
    bold_chars = 0
    for run in paragraph.runs:
        run_text = run.text or ""
        total_chars += len(run_text)
        if run.bold:
            bold_chars += len(run_text)
    if total_chars == 0:
        return False
    return bold_chars / total_chars >= 0.8


def extract_notes_from_document(doc: DocumentType) -> List[Tuple[str, str]]:  # type: ignore[valid-type]
    notes: List[Tuple[str, str]] = []
    current_title: Optional[str] = None
    current_body: List[str] = []
    for paragraph in doc.paragraphs:
        raw_text = paragraph.text or ""
        text = raw_text.strip()
        if not text:
            continue
        if is_paragraph_bold(paragraph):
            if current_title is not None:
                notes.append((current_title, "\n".join(current_body).strip()))
            current_title = text
            current_body = []
        else:
            if current_title is None:
                # Ignoramos texto suelto antes del primer título
                continue
            current_body.append(text)
    if current_title is not None:
        notes.append((current_title, "\n".join(current_body).strip()))
    return notes


def load_notes_from_docx(root: str) -> Tuple[List[Note], Dict[str, List[str]]]:
    if not DOCX_AVAILABLE or Document is None:  # pragma: no cover - fallback
        raise RuntimeError("python-docx no está disponible; instalalo o usá --notes-report")

    notes: List[Note] = []
    images_by_page: Dict[str, List[str]] = {}
    for entry in sorted(os.listdir(root)):
        folder_path = os.path.join(root, entry)
        if not os.path.isdir(folder_path):
            continue
        page_key = normalize_page_key(entry)
        docx_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx")]
        image_files = [
            f
            for f in os.listdir(folder_path)
            if os.path.splitext(f)[1].lower() in {".jpg", ".jpeg", ".png", ".tif", ".tiff"}
        ]
        images_by_page[page_key] = image_files
        for docx_name in sorted(docx_files):
            doc_path = os.path.join(folder_path, docx_name)
            try:
                doc = Document(doc_path)  # type: ignore[call-arg]
            except Exception as exc:  # pylint: disable=broad-except
                print(f"No se pudo abrir '{doc_path}': {exc}", file=sys.stderr)
                continue
            raw_notes = extract_notes_from_document(doc)
            if not raw_notes:
                # Una nota sin división: usamos todo el cuerpo del documento
                text = "\n".join(p.text for p in doc.paragraphs if p.text)
                raw_notes = [(docx_name, text)]
            for idx, (title, body) in enumerate(raw_notes, start=1):
                note = Note(
                    page_hint=page_key,
                    docx_file=docx_name,
                    note_index=idx,
                    title=title.strip(),
                    body_text=body.strip(),
                    title_chars=len(title.strip()),
                    body_chars=len(body.replace("\n", "").strip()),
                    image_count=0,
                )
                notes.append(note)
    return notes, images_by_page


def load_notes_from_report(path: str) -> Tuple[List[Note], Dict[str, List[str]]]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        notes: List[Note] = []
        images_by_page: Dict[str, List[str]] = defaultdict(list)
        with open(path, "r", encoding="utf-8-sig") as fh:
            reader = csv.DictReader(fh)
            for row in reader:
                page_folder = normalize_page_key(row.get("page_folder", ""))
                docx_file = row.get("docx_file") or None
                title = (row.get("title") or "").strip()
                note_index = int(row.get("note_index", 1))
                body_chars = int(row.get("char_count", 0))
                images_raw = row.get("images_list", "")
                image_files = [part for part in images_raw.split(";") if part]
                if image_files:
                    images_by_page[page_folder].extend(image_files)
                note = Note(
                    page_hint=page_folder,
                    docx_file=docx_file,
                    note_index=note_index,
                    title=title,
                    body_text="",
                    title_chars=len(title),
                    body_chars=body_chars,
                    image_count=int(row.get("images_count", 0) or 0),
                )
                notes.append(note)
        return notes, images_by_page

    if ext == ".json":
        with open(path, "r", encoding="utf-8") as fh:
            payload = json.load(fh)
        notes: List[Note] = []
        images_by_page: Dict[str, List[str]] = defaultdict(list)
        pages = payload.get("pages", {}) if isinstance(payload, dict) else {}
        for page_name, info in pages.items():
            page_key = normalize_page_key(page_name)
            image_files = info.get("images", []) if isinstance(info, dict) else []
            if image_files:
                images_by_page[page_key].extend(image_files)
            docx_files = info.get("docx", []) if isinstance(info, dict) else []
            note_count = int(info.get("notes", 0)) if isinstance(info, dict) else 0
            total_chars = int(info.get("chars", 0)) if isinstance(info, dict) else 0
            # Distribuimos caracteres promedio si no tenemos detalle
            avg_chars = total_chars // note_count if note_count else 0
            for idx in range(1, note_count + 1):
                title = f"Nota {idx}"
                note = Note(
                    page_hint=page_key,
                    docx_file=docx_files[0] if docx_files else None,
                    note_index=idx,
                    title=title,
                    body_text="",
                    title_chars=len(title),
                    body_chars=avg_chars,
                    image_count=len(image_files) // max(1, note_count),
                )
                notes.append(note)
        return notes, images_by_page

    raise ValueError("Formato de --notes-report no soportado; use CSV o JSON")


def detect_column_width(slots: Sequence[Slot], default_width_pt: float = 124.0) -> float:
    widths = [round(slot.width_pt, 2) for slot in slots if slot.width_pt > 30]
    if not widths:
        return default_width_pt
    freq = Counter(widths)
    most_common_width, _ = freq.most_common(1)[0]
    return most_common_width


def compute_text_height(
    chars: int,
    width_pt: float,
    font_size_pt: float,
    column_model: ColumnModel,
) -> float:
    if chars <= 0:
        return 0.0
    chars_per_line = max(column_model.chars_per_line(width_pt, font_size_pt), 1.0)
    lines = math.ceil(chars / chars_per_line)
    return lines * column_model.line_height(font_size_pt)


def compute_image_height(slot_width_pt: float, config: Dict[str, object]) -> float:
    image_cfg = config.get("image", {}) if isinstance(config, dict) else {}
    width_cm = float(image_cfg.get("width_cm", IMAGE_DEFAULT_WIDTH_CM))
    min_height_cm = float(image_cfg.get("min_height_cm", IMAGE_MIN_HEIGHT_CM))
    width_pt = width_cm * POINTS_PER_CM
    height_pt = min_height_cm * POINTS_PER_CM
    if slot_width_pt < width_pt:
        # Ajustamos proporcionalmente manteniendo la relación de aspecto aproximada
        scale = slot_width_pt / max(width_pt, 1.0)
        height_pt = height_pt * scale
    return height_pt


def compute_note_height(
    note: Note,
    slot: Slot,
    font_title: float,
    font_body: float,
    column_model: ColumnModel,
    config: Dict[str, object],
) -> Tuple[float, Dict[str, float]]:
    title_cfg = config.get("title", {}) if isinstance(config, dict) else {}
    body_cfg = config.get("body", {}) if isinstance(config, dict) else {}
    image_cfg = config.get("image", {}) if isinstance(config, dict) else {}

    title_spacing = float(title_cfg.get("spacing_after_pt", 8.0)) if note.title_chars else 0.0
    body_spacing = float(body_cfg.get("spacing_after_pt", 6.0)) if note.body_chars else 0.0
    image_spacing = float(image_cfg.get("spacing_pt", 6.0)) if note.has_image() else 0.0

    title_height = compute_text_height(note.title_chars, slot.width_pt, font_title, column_model)
    body_height = compute_text_height(note.body_chars, slot.width_pt, font_body, column_model)
    image_height = compute_image_height(slot.width_pt, config) if note.has_image() else 0.0

    heights = {
        "title_pt": title_height,
        "body_pt": body_height,
        "image_pt": image_height,
        "title_spacing_pt": title_spacing,
        "body_spacing_pt": body_spacing,
        "image_spacing_pt": image_spacing,
    }

    total = title_height + body_height + image_height + title_spacing + body_spacing + image_spacing
    return total, heights


def compute_capacity(
    available_height: float,
    slot_width_pt: float,
    font_body: float,
    column_model: ColumnModel,
    config: Dict[str, object],
) -> int:
    body_cfg = config.get("body", {}) if isinstance(config, dict) else {}
    body_spacing = float(body_cfg.get("spacing_after_pt", 6.0))
    chars_per_line = column_model.chars_per_line(slot_width_pt, font_body)
    line_height = column_model.line_height(font_body)
    effective_height = max(available_height - body_spacing, 0)
    lines = math.floor(effective_height / line_height)
    return int(lines * chars_per_line)


def try_note_in_slot(
    note: Note,
    slot: Slot,
    remaining_height: float,
    column_model: ColumnModel,
    config: Dict[str, object],
) -> LayoutResult:
    title_cfg = config.get("title", {}) if isinstance(config, dict) else {}
    body_cfg = config.get("body", {}) if isinstance(config, dict) else {}

    title_sizes = sorted(
        {float(title_cfg.get("pt_base", 25.0)), float(title_cfg.get("pt_min", 24.0)), float(title_cfg.get("pt_max", 26.0))}
    )
    body_sizes = sorted(
        {float(body_cfg.get("pt_base", 9.5)), float(body_cfg.get("pt_min", 9.0)), float(body_cfg.get("pt_max", 10.0))}
    )

    attempt_history: List[Dict[str, object]] = []
    best_fit: Optional[LayoutResult] = None
    minimal_overflow = math.inf

    for body_size in sorted(body_sizes, reverse=True):
        for title_size in sorted(title_sizes, reverse=True):
            total_height, heights = compute_note_height(note, slot, title_size, body_size, column_model, config)
            attempt_history.append(
                {
                    "title_pt": title_size,
                    "body_pt": body_size,
                    "total_height_pt": total_height,
                }
            )
            remaining = remaining_height - total_height
            if remaining >= 0:
                result = LayoutResult(
                    note=note,
                    slot=slot,
                    fit=True,
                    font_title=title_size,
                    font_body=body_size,
                    heights_pt=heights,
                    remaining_height_pt=remaining,
                    overflow_chars=0,
                    used_height_pt=total_height,
                    attempt_summary={"tested": attempt_history},
                )
                return result

            capacity = compute_capacity(remaining_height, slot.width_pt, body_size, column_model, config)
            overflow = max(note.body_chars - capacity, 0)
            missing_height = max(total_height - remaining_height, 0.0)
            if overflow < minimal_overflow:
                minimal_overflow = overflow
                best_fit = LayoutResult(
                    note=note,
                    slot=slot,
                    fit=False,
                    font_title=title_size,
                    font_body=body_size,
                    heights_pt=heights,
                    remaining_height_pt=remaining,
                    overflow_chars=overflow,
                    used_height_pt=total_height,
                    attempt_summary={
                        "tested": attempt_history.copy(),
                        "missing_height_pt": missing_height,
                    },
                )

    if best_fit is None:
        best_fit = LayoutResult(
            note=note,
            slot=slot,
            fit=False,
            font_title=float(title_cfg.get("pt_min", 24.0)),
            font_body=float(body_cfg.get("pt_min", 9.0)),
            heights_pt={"title_pt": 0.0, "body_pt": 0.0, "image_pt": 0.0},
            remaining_height_pt=remaining_height,
            overflow_chars=note.body_chars,
            used_height_pt=0.0,
            attempt_summary={"tested": attempt_history},
        )
    return best_fit


def plan_layout(
    slots: Sequence[Slot],
    notes: Sequence[Note],
    images_by_page: Dict[str, List[str]],
    config: Dict[str, object],
    image_strategy: str,
) -> Dict[str, object]:
    slots_by_page: Dict[str, List[Slot]] = defaultdict(list)
    for slot in slots:
        page_key = normalize_page_key(slot.page_name or str(slot.page_index))
        slots_by_page[page_key].append(slot)

    for page_slots in slots_by_page.values():
        page_slots.sort(key=lambda s: (s.y_pt, s.x_pt))

    base_column_width_pt = detect_column_width(slots)
    column_model = ColumnModel(
        column_width_pt=base_column_width_pt,
        gutter_pt=float(config.get("column", {}).get("gutter_pt", 12.0)) if isinstance(config.get("column"), dict) else 12.0,
        char_width_factor=float(config.get("body", {}).get("char_width_factor", 0.42))
        if isinstance(config.get("body"), dict)
        else 0.42,
        leading_factor=float(config.get("body", {}).get("leading_factor", 1.2))
        if isinstance(config.get("body"), dict)
        else 1.2,
    )

    grouped_notes: Dict[str, List[Note]] = defaultdict(list)
    for note in notes:
        key = normalize_page_key(note.page_hint)
        grouped_notes[key].append(note)

    placements_by_page: Dict[str, List[LayoutResult]] = defaultdict(list)
    unplaced_by_page: Dict[str, List[UnplacedNote]] = defaultdict(list)

    alias_mapping: Dict[str, str] = {}

    for original_key, page_notes in grouped_notes.items():
        page_key = original_key
        slots_for_page = slots_by_page.get(page_key)
        if not slots_for_page:
            alt_key = None
            idx = first_int_in_string(page_key)
            if idx is not None:
                alt_key = normalize_page_key(str(idx))
            if alt_key and slots_by_page.get(alt_key):
                slots_for_page = slots_by_page[alt_key]
                alias_mapping[original_key] = alt_key
                page_key = alt_key
            else:
                unplaced_by_page[original_key].extend(
                    UnplacedNote(note=n, reason="no_slot_for_page") for n in page_notes
                )
                continue

        available_images = images_by_page.get(original_key, []) or images_by_page.get(page_key, [])
        assign_images(page_notes, available_images, image_strategy)

        slot_iter = iter(slots_for_page)
        current_slot = next(slot_iter, None)
        remaining_height = current_slot.height_pt if current_slot else 0.0

        for note in page_notes:
            placed_result: Optional[LayoutResult] = None
            last_attempt: Optional[LayoutResult] = None
            while current_slot:
                result = try_note_in_slot(note, current_slot, remaining_height, column_model, config)
                if result.fit:
                    placed_result = result
                    break
                last_attempt = result
                current_slot = next(slot_iter, None)
                remaining_height = current_slot.height_pt if current_slot else 0.0

            if placed_result:
                placements_by_page[current_slot.page_name].append(placed_result)  # type: ignore[arg-type]
                remaining_height = placed_result.remaining_height_pt - float(config.get("note_spacing_pt", 10.0))
                if remaining_height <= 0:
                    current_slot = next(slot_iter, None)
                    remaining_height = current_slot.height_pt if current_slot else 0.0
                continue

            # No se pudo ubicar en ningún slot restante
            if last_attempt is not None:
                unplaced_by_page[page_key].append(
                    UnplacedNote(
                        note=note,
                        reason="overflow",
                        overflow_chars=last_attempt.overflow_chars,
                        missing_height_pt=max(-last_attempt.remaining_height_pt, 0.0),
                    )
                )
            else:
                unplaced_by_page[page_key].append(UnplacedNote(note=note, reason="no_slots_left"))

    pages_output: List[Dict[str, object]] = []
    for page_key, slot_list in slots_by_page.items():
        placements = placements_by_page.get(page_key, [])
        unplaced = unplaced_by_page.get(page_key, [])
        remaining_height_by_slot: Dict[str, float] = {slot.slot_id: slot.height_pt for slot in slot_list}
        for placement in placements:
            remaining_height_by_slot[placement.slot.slot_id] = max(
                remaining_height_by_slot.get(placement.slot.slot_id, placement.slot.height_pt) - placement.used_height_pt, 0.0
            )
        pages_output.append(
            {
                "page_name": page_key,
                "page_index": slot_list[0].page_index if slot_list else 0,
                "slots": [slot.to_dict() for slot in slot_list],
                "placements": [placement.to_dict() for placement in placements],
                "remaining_height_pt": remaining_height_by_slot,
                "unplaced_notes": [entry.to_dict() for entry in unplaced],
            }
        )

    # Notas de páginas sin slots (quedaron en el dict)
    for page_key, entries in unplaced_by_page.items():
        if page_key in slots_by_page:
            continue
        pages_output.append(
            {
                "page_name": page_key,
                "page_index": None,
                "slots": [],
                "placements": [],
                "remaining_height_pt": {},
                "unplaced_notes": [entry.to_dict() for entry in entries],
            }
        )

    output = {
        "config": config,
        "column_model": {
            "base_width_pt": column_model.column_width_pt,
            "gutter_pt": column_model.gutter_pt,
            "char_width_factor": column_model.char_width_factor,
            "leading_factor": column_model.leading_factor,
        },
        "page_aliases": alias_mapping,
        "pages": pages_output,
    }

    return output


def write_json_output(path: str, payload: Dict[str, object]) -> None:
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, indent=2, ensure_ascii=False)


def write_csv_output(path: str, payload: Dict[str, object]) -> None:
    rows: List[Dict[str, object]] = []
    for page in payload.get("pages", []):
        page_name = page.get("page_name", "")
        for placement in page.get("placements", []):
            rows.append(
                {
                    "page": page_name,
                    "note_id": placement.get("note_id"),
                    "slot_id": placement.get("slot_id"),
                    "fit": placement.get("fit"),
                    "overflow_chars": placement.get("overflow_chars"),
                    "font_title_pt": placement.get("font_title_pt"),
                    "font_body_pt": placement.get("font_body_pt"),
                    "used_height_pt": placement.get("used_height_pt"),
                    "remaining_height_pt": placement.get("remaining_height_pt"),
                }
            )
        for unplaced in page.get("unplaced_notes", []):
            rows.append(
                {
                    "page": page_name,
                    "note_id": unplaced.get("note_id"),
                    "slot_id": None,
                    "fit": False,
                    "overflow_chars": unplaced.get("overflow_chars"),
                    "font_title_pt": None,
                    "font_body_pt": None,
                    "used_height_pt": None,
                    "remaining_height_pt": unplaced.get("missing_height_pt"),
                }
            )

    fieldnames = [
        "page",
        "note_id",
        "slot_id",
        "fit",
        "overflow_chars",
        "font_title_pt",
        "font_body_pt",
        "used_height_pt",
        "remaining_height_pt",
    ]
    with open(path, "w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def pretty_print(payload: Dict[str, object]) -> None:
    pages = payload.get("pages", [])
    print("Resumen de planificación")
    for page in pages:
        name = page.get("page_name", "")
        placements = page.get("placements", [])
        unplaced = page.get("unplaced_notes", [])
        print(f"\nPágina {name}: {len(placements)} notas ubicadas, {len(unplaced)} pendientes")
        for placement in placements:
            status = "OK" if placement.get("fit") else "OVERFLOW"
            note_title = placement.get("title", placement.get("note_id"))
            overflow = placement.get("overflow_chars", 0)
            remaining = placement.get("remaining_height_pt", 0.0)
            print(
                f"  - {status} · {note_title} · overflow={overflow} chars · rem={remaining:.1f} pt"
            )
        for un in unplaced:
            note_title = un.get("title", un.get("note_id"))
            reason = un.get("reason")
            print(f"  - PENDIENTE · {note_title} · motivo: {reason}")


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_arguments(argv)
    config = load_config(args.config)

    slots: List[Slot] = []

    if not args.slots_json and not args.idml:
        try:
            user_input = input(
                "Ingresá la ruta del archivo de slots (JSON) o del documento IDML: "
            ).strip()
        except EOFError:
            user_input = ""

        user_input = os.path.expanduser(user_input.strip('"').strip("'"))

        if user_input.lower().endswith(".json"):
            args.slots_json = user_input
        elif user_input:
            args.idml = user_input

    if args.slots_json:
        slots = load_slots_from_json(args.slots_json)
    elif args.idml:
        slots = load_slots_from_idml(args.idml)
    else:
        print("Debés indicar --slots-json o --idml", file=sys.stderr)
        return 2

    if not slots:
        print("No se encontraron slots válidos", file=sys.stderr)
        return 3

    notes: List[Note] = []
    images_by_page: Dict[str, List[str]] = defaultdict(list)

    if args.notes_root:
        try:
            notes, images_by_page = load_notes_from_docx(args.notes_root)
        except Exception as exc:  # pylint: disable=broad-except
            print(f"No se pudieron cargar DOCX: {exc}", file=sys.stderr)
            return 4
    elif args.notes_report:
        notes, images_by_page = load_notes_from_report(args.notes_report)
    else:
        print("Debés indicar --notes-root o --notes-report", file=sys.stderr)
        return 5

    if not notes:
        print("No se detectaron notas", file=sys.stderr)
        return 6

    payload = plan_layout(slots, notes, images_by_page, config, args.default_image_assignment)
    write_json_output(args.output_json, payload)
    if args.output_csv:
        write_csv_output(args.output_csv, payload)
    if args.pretty:
        pretty_print(payload)

    return 0


if __name__ == "__main__":
    sys.exit(main())

